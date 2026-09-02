from __future__ import annotations

import argparse
import csv
import glob
import os
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


# Default root is 2 levels up from 'формат для печати/tools'
DEFAULT_ROOT = Path(__file__).resolve().parents[2]
ROOT = Path(os.environ.get('BOOK_ROOT', str(DEFAULT_ROOT)))
OUT = ROOT / 'формат для печати' / '_единая рукопись.docx'
CHANGELOG = ROOT / 'формат для печати' / 'журнал форматных изменений.csv'
EXCEPTIONS = ROOT / 'формат для печати' / 'исключения для проверки.md'
TOC_MAP = Path('/private/tmp/book2-toc-pages.json')

NS = {
    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
    'style': 'urn:oasis:names:tc:opendocument:xmlns:style:1.0',
    'fo': 'urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0',
}
P = f'{{{NS["text"]}}}p'
S = f'{{{NS["text"]}}}s'
SPAN = f'{{{NS["text"]}}}span'
LB = f'{{{NS["text"]}}}line-break'
TAB = f'{{{NS["text"]}}}tab'
STYLE = f'{{{NS["style"]}}}style'
TEXT_PROPS = f'{{{NS["style"]}}}text-properties'

PART_TITLES = {
    1: 'I. Введение',
    2: 'II. Город',
    3: 'III. Коротко',
    4: 'IV. Личное',
    5: 'V. Общество',
}


@dataclass(frozen=True)
class CharStyle:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    superscript: bool = False
    subscript: bool = False


@dataclass
class SourceParagraph:
    source: str
    number: int
    runs: list[tuple[str, CharStyle]]
    text: str


class Logger:
    def __init__(self) -> None:
        self.rows: list[dict[str, str]] = []
        self.exceptions: list[str] = []

    def change(self, source: str, paragraph: int, rule: str, old: str, new: str) -> None:
        self.rows.append({
            'source_file': source,
            'paragraph': str(paragraph),
            'rule': rule,
            'old_value': old.replace('\n', '↵'),
            'new_value': new.replace('\n', '↵'),
        })


LOG = Logger()


def attr(el: ET.Element, namespace: str, name: str) -> str | None:
    return el.attrib.get(f'{{{namespace}}}{name}')


def style_map(root: ET.Element) -> dict[str, CharStyle]:
    styles: dict[str, CharStyle] = {}
    for el in root.iter(STYLE):
        name = attr(el, NS['style'], 'name')
        if not name:
            continue
        props = next((x for x in el if x.tag == TEXT_PROPS), None)
        if props is None:
            continue
        styles[name] = CharStyle(
            bold=attr(props, NS['fo'], 'font-weight') in {'bold', '700', '800', '900'},
            italic=attr(props, NS['fo'], 'font-style') == 'italic',
            underline=attr(props, NS['style'], 'text-underline-style') not in {None, 'none'},
            superscript=attr(props, NS['style'], 'text-position') == 'super',
            subscript=attr(props, NS['style'], 'text-position') == 'sub',
        )
    return styles


def merge_style(base: CharStyle, extra: CharStyle | None) -> CharStyle:
    if extra is None:
        return base
    return CharStyle(
        bold=base.bold or extra.bold,
        italic=base.italic or extra.italic,
        underline=base.underline or extra.underline,
        superscript=base.superscript or extra.superscript,
        subscript=base.subscript or extra.subscript,
    )


def read_odt(path: Path) -> list[SourceParagraph]:
    with zipfile.ZipFile(path) as zf:
        content = ET.fromstring(zf.read('content.xml'))
        styles_xml = ET.fromstring(zf.read('styles.xml'))
    styles = style_map(content)
    styles.update(style_map(styles_xml))
    rel = str(path.relative_to(ROOT))
    out: list[SourceParagraph] = []

    def walk(node: ET.Element, inherited: CharStyle, runs: list[tuple[str, CharStyle]]) -> None:
        if node.text:
            runs.append((node.text, inherited))
        for child in node:
            current = inherited
            if child.tag == SPAN:
                current = merge_style(inherited, styles.get(attr(child, NS['text'], 'style-name') or ''))
                walk(child, current, runs)
            elif child.tag == S:
                runs.append((' ' * int(attr(child, NS['text'], 'c') or '1'), inherited))
            elif child.tag == LB:
                runs.append(('\n', inherited))
            elif child.tag == TAB:
                runs.append(('\t', inherited))
            else:
                walk(child, current, runs)
            if child.tail:
                runs.append((child.tail, inherited))

    for number, para in enumerate(content.iter(P), start=1):
        runs: list[tuple[str, CharStyle]] = []
        walk(para, CharStyle(), runs)
        out.append(SourceParagraph(rel, number, runs, ''.join(t for t, _ in runs)))
    return out


def normalize_runs(para: SourceParagraph) -> list[tuple[str, CharStyle]]:
    chars: list[tuple[str, CharStyle]] = []
    for text, sty in para.runs:
        chars.extend((c, sty) for c in text)
    before = ''.join(c for c, _ in chars)
    # Tabs and repeated spaces are typewriter layout. Manual line breaks remain intact.
    collapsed: list[tuple[str, CharStyle]] = []
    seen_space = False
    for char, sty in chars:
        if char == '\t':
            char = ' '
        if char == ' ':
            if seen_space:
                continue
            seen_space = True
        else:
            seen_space = False
        collapsed.append((char, sty))
    while collapsed and collapsed[0][0] == ' ':
        collapsed.pop(0)
    while collapsed and collapsed[-1][0] == ' ':
        collapsed.pop()

    # Approved typographic normalization: ... → …
    ellipsis: list[tuple[str, CharStyle]] = []
    i = 0
    while i < len(collapsed):
        if ''.join(c for c, _ in collapsed[i:i + 3]) == '...':
            ellipsis.append(('…', collapsed[i][1]))
            i += 3
        else:
            ellipsis.append(collapsed[i])
            i += 1

    # Approved date normalization: retain content, remove only accidental inner spaces.
    current = ''.join(c for c, _ in ellipsis)
    normalized = re.sub(r'\b(\d{1,2})\.\s+(\d{1,2})\.\s+(\d{4})\b', r'\1.\2.\3', current)
    if normalized != current:
        rebuilt: list[tuple[str, CharStyle]] = []
        pos = 0
        for match in re.finditer(r'\b(\d{1,2})\.\s+(\d{1,2})\.\s+(\d{4})\b', current):
            rebuilt.extend(ellipsis[pos:match.start()])
            rebuilt.extend((c, ellipsis[match.start()][1]) for c in re.sub(r'\s+', '', match.group()))
            pos = match.end()
        rebuilt.extend(ellipsis[pos:])
        ellipsis = rebuilt

    after = ''.join(c for c, _ in ellipsis)
    if before != after:
        rules = []
        if '\t' in before or re.search(r'(^ | $| {2,})', before):
            rules.append('whitespace')
        if '...' in before:
            rules.append('ellipsis')
        if re.search(r'\b\d{1,2}\.\s+\d{1,2}\.\s+\d{4}\b', before):
            rules.append('date-spacing')
        LOG.change(para.source, para.number, '+'.join(rules) or 'normalization', before, after)

    out: list[tuple[str, CharStyle]] = []
    for char, sty in ellipsis:
        if out and out[-1][1] == sty:
            out[-1] = (out[-1][0] + char, sty)
        else:
            out.append((char, sty))
    return out


def text_of(runs: list[tuple[str, CharStyle]]) -> str:
    return ''.join(t for t, _ in runs)


def has_initial_bold(runs: list[tuple[str, CharStyle]]) -> bool:
    for text, sty in runs:
        if text.strip():
            return sty.bold
    return False


def is_upperish(text: str) -> bool:
    letters = [c for c in text if c.isalpha()]
    return len(letters) >= 3 and sum(c.isupper() for c in letters) / len(letters) >= 0.8


def candidate_title(text: str) -> bool:
    value = text.strip()
    if not value or value.startswith('***') or len(value) > 115:
        return False
    if value.startswith(('-', '—')):
        return False
    return is_upperish(value) or (len(value) <= 65 and not value.endswith(('.', ';')))


def classify_chapter(paras: list[SourceParagraph]) -> tuple[int | None, set[int], set[int], list[str]]:
    """Return title index, subtitle/front-note indices, opening-epigraph indices, and review notes."""
    usable = [(i, p) for i, p in enumerate(paras) if p.text.strip()]
    notes: list[str] = []
    if not usable:
        return None, set(), set(), ['Empty file.']

    # 1. Determine title
    title: int | None = None
    first_idx, first_p = usable[0]
    if candidate_title(first_p.text):
        title = first_idx
    else:
        # Check within the first few non-empty paragraphs for an upperish / strong title
        for idx, p in usable[:6]:
            val = p.text.strip()
            if is_upperish(val) and not val.startswith('***') and len(val) <= 120:
                title = idx
                break

    extras: set[int] = set()
    epigraph: set[int] = set()

    if title is None:
        notes.append('Не удалось уверенно распознать заголовок главы (текст начинается без явного заголовка); файл начат с новой страницы.')
        return None, extras, epigraph, notes

    # Find position of title in usable list
    title_pos = next(pos for pos, (idx, _) in enumerate(usable) if idx == title)

    # Check elements before title (epigraphs, front notes, intro labels)
    for idx, p in usable[:title_pos]:
        runs = normalize_runs(p)
        if runs and all(sty.italic for text, sty in runs if text.strip()):
            epigraph.add(idx)
        else:
            extras.add(idx)

    # Check immediate subtitle after title
    if title_pos + 1 < len(usable):
        next_idx, next_p = usable[title_pos + 1]
        next_val = next_p.text.strip()
        if candidate_title(next_val) and len(next_val) <= 100 and not next_val.startswith('***'):
            extras.add(next_idx)

    return title, extras, epigraph, notes


def set_font(run, name: str = 'Georgia', size: float | None = None) -> None:
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rfonts.set(qn(f'w:{key}'), name)
    if size is not None:
        run.font.size = Pt(size)


def define_styles(doc: Document) -> None:
    normal = doc.styles['Normal']
    normal.font.name = 'Georgia'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Georgia')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Georgia')
    normal._element.rPr.rFonts.set(qn('w:cs'), 'Georgia')
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Mm(8.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    def custom(name: str, base: str = 'Normal'):
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, builtin=False) if name not in doc.styles else doc.styles[name]

    part = doc.styles['Heading 1']
    part.base_style = normal
    part.font.name = 'Georgia'; part.font.size = Pt(16); part.font.bold = True; part.font.color.rgb = RGBColor(0, 0, 0)
    part.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part.paragraph_format.first_line_indent = Mm(0)
    part.paragraph_format.space_before = Pt(0); part.paragraph_format.space_after = Pt(18)
    part.paragraph_format.keep_with_next = True

    chapter = doc.styles['Heading 2']
    chapter.base_style = normal
    chapter.font.name = 'Georgia'; chapter.font.size = Pt(13); chapter.font.bold = True; chapter.font.color.rgb = RGBColor(0, 0, 0)
    chapter.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter.paragraph_format.first_line_indent = Mm(0)
    chapter.paragraph_format.space_before = Pt(0); chapter.paragraph_format.space_after = Pt(14)
    chapter.paragraph_format.keep_with_next = True
    chapter.paragraph_format.page_break_before = True

    subtitle = custom('Chapter Subtitle')
    subtitle.base_style = normal
    subtitle.font.name = 'Georgia'; subtitle.font.size = Pt(11); subtitle.font.italic = True
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Mm(0)
    subtitle.paragraph_format.space_before = Pt(0); subtitle.paragraph_format.space_after = Pt(14)

    front = custom('Front Note')
    front.base_style = normal
    front.font.name = 'Georgia'; front.font.size = Pt(10)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.first_line_indent = Mm(0)
    front.paragraph_format.space_before = Pt(0); front.paragraph_format.space_after = Pt(4)

    epigraph = custom('Opening Epigraph')
    epigraph.base_style = normal
    epigraph.font.name = 'Georgia'; epigraph.font.size = Pt(10); epigraph.font.italic = True
    epigraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    epigraph.paragraph_format.left_indent = Mm(50)
    epigraph.paragraph_format.first_line_indent = Mm(0)
    epigraph.paragraph_format.space_before = Pt(0); epigraph.paragraph_format.space_after = Pt(2)

    source = custom('Epigraph Source')
    source.base_style = normal
    source.font.name = 'Georgia'; source.font.size = Pt(10)
    source.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    source.paragraph_format.left_indent = Mm(50)
    source.paragraph_format.first_line_indent = Mm(0)
    source.paragraph_format.space_before = Pt(0); source.paragraph_format.space_after = Pt(12)

    quote = custom('Block Quote')
    quote.base_style = normal
    quote.font.name = 'Georgia'; quote.font.size = Pt(10.5)
    quote.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    quote.paragraph_format.left_indent = Mm(8); quote.paragraph_format.right_indent = Mm(8)
    quote.paragraph_format.first_line_indent = Mm(0)
    quote.paragraph_format.space_before = Pt(6); quote.paragraph_format.space_after = Pt(6)

    separator = custom('Topic Separator')
    separator.base_style = normal
    separator.font.name = 'Georgia'; separator.font.size = Pt(11)
    separator.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.paragraph_format.first_line_indent = Mm(0)
    separator.paragraph_format.space_before = Pt(8); separator.paragraph_format.space_after = Pt(8)

    dash_list = custom('Dash List')
    dash_list.base_style = normal
    dash_list.font.name = 'Georgia'; dash_list.font.size = Pt(11)
    dash_list.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dash_list.paragraph_format.left_indent = Mm(8.5); dash_list.paragraph_format.first_line_indent = Mm(-5)
    dash_list.paragraph_format.space_before = Pt(0); dash_list.paragraph_format.space_after = Pt(0)
    dash_list.paragraph_format.line_spacing = 1.0

    cover = custom('Cover Text')
    cover.base_style = normal
    cover.font.name = 'Georgia'; cover.font.size = Pt(25)
    cover.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.first_line_indent = Mm(0)
    cover.paragraph_format.space_before = Pt(0); cover.paragraph_format.space_after = Pt(6)

    toc = custom('TOC Title')
    toc.base_style = normal
    toc.font.name = 'Georgia'; toc.font.size = Pt(14); toc.font.bold = True
    toc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc.paragraph_format.first_line_indent = Mm(0)
    toc.paragraph_format.space_before = Pt(0); toc.paragraph_format.space_after = Pt(12)

    toc_part = custom('TOC Part')
    toc_part.base_style = normal
    toc_part.font.name = 'Georgia'; toc_part.font.size = Pt(12); toc_part.font.bold = True
    toc_part.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_part.paragraph_format.first_line_indent = Mm(0)
    toc_part.paragraph_format.space_before = Pt(5); toc_part.paragraph_format.space_after = Pt(1)

    toc_chapter = custom('TOC Chapter')
    toc_chapter.base_style = normal
    toc_chapter.font.name = 'Georgia'; toc_chapter.font.size = Pt(9)
    toc_chapter.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_chapter.paragraph_format.left_indent = Mm(5)
    toc_chapter.paragraph_format.first_line_indent = Mm(0)
    toc_chapter.paragraph_format.space_before = Pt(0); toc_chapter.paragraph_format.space_after = Pt(0)


def configure_section(section) -> None:
    section.page_width = Mm(148)
    section.page_height = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(9)


def add_page_number(footer) -> None:
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    set_font(run, size=9)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.addnext(fld)


def set_page_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn('w:pgNumType'))
    if pg is None:
        pg = OxmlElement('w:pgNumType')
        sect_pr.append(pg)
    pg.set(qn('w:start'), str(start))


def continue_page_numbers(section) -> None:
    """Do not copy the previous section's explicit page-number restart."""
    pg = section._sectPr.find(qn('w:pgNumType'))
    if pg is not None:
        section._sectPr.remove(pg)


def toc_entry_details() -> list[tuple[str, str, str, str]]:
    """Return kind, printed title, key, and source phrase used to locate its rendered page."""
    entries: list[tuple[str, str, str, str]] = []
    current_part = None
    for part, _chapter, path in document_files():
        if part != current_part:
            part_title = PART_TITLES.get(part, f'{part}. ЧАСТЬ')
            entries.append(('part', part_title, f'part:{part}', part_title))
            current_part = part
        paragraphs = read_odt(path)
        title, _extras, _epigraphs, _notes = classify_chapter(paragraphs)
        if title is not None:
            title_text = re.sub(r'\s+', ' ', paragraphs[title].text.strip()).replace('...', '…')
            search_text = title_text
        else:
            title_text = re.sub(r'^\d+-\d{2}\s+', '', path.stem)
            first = next((p.text.strip() for p in paragraphs if p.text.strip()), '')
            search_text = re.sub(r'^\*\*\*\s*', '', first)
        entries.append(('chapter', title_text, f'chapter:{path.relative_to(ROOT)}', search_text))
    return entries


def toc_entries() -> list[tuple[str, str, str]]:
    return [(kind, title, key) for kind, title, key, _search in toc_entry_details()]


def add_static_toc(doc: Document, entries: list[tuple[str, str, str]]) -> None:
    page_map: dict[str, int] = {}
    if TOC_MAP.exists():
        import json
        page_map = json.loads(TOC_MAP.read_text(encoding='utf-8'))
    doc.add_paragraph('СОДЕРЖАНИЕ', style='TOC Title')
    for kind, title, key in entries:
        p = doc.add_paragraph(style='TOC Part' if kind == 'part' else 'TOC Chapter')
        p.paragraph_format.tab_stops.add_tab_stop(Mm(112.7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(title + '\t' + str(page_map.get(key, '—')))
        set_font(run, size=10 if kind == 'part' else 9)
        if kind == 'part':
            run.bold = True
    doc.add_page_break()


def add_runs(paragraph, runs: list[tuple[str, CharStyle]]) -> None:
    for text, sty in runs:
        parts = text.split('\n')
        for idx, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                set_font(run)
                if sty.bold:
                    run.bold = True
                if sty.italic:
                    run.italic = True
                if sty.underline:
                    run.underline = True
                if sty.superscript:
                    run.font.superscript = True
                if sty.subscript:
                    run.font.subscript = True
            if idx < len(parts) - 1:
                paragraph.add_run().add_break()


def add_normalized_paragraph(doc: Document, para: SourceParagraph, style: str = 'Normal', override_runs=None) -> None:
    runs = override_runs if override_runs is not None else normalize_runs(para)
    p = doc.add_paragraph(style=style)
    add_runs(p, runs)


def document_files() -> list[tuple[int, int, Path]]:
    result: list[tuple[int, int, Path]] = []
    for path_s in glob.glob(str(ROOT / '*.odt')):
        path = Path(path_s)
        match = re.match(r'^([1-9]\d*)-(\d{2})\s+', path.name)
        if match:
            result.append((int(match.group(1)), int(match.group(2)), path))
    return sorted(result)


def is_list_run(items: list[tuple[int, SourceParagraph]]) -> set[int]:
    result: set[int] = set()
    run: list[int] = []
    for index, para in items:
        value = text_of(normalize_runs(para)).lstrip()
        if re.match(r'^-\s+', value) or re.match(r'^\d+[.)]\s+', value):
            run.append(index)
        else:
            if len(run) >= 2:
                result.update(run)
            run = []
    if len(run) >= 2:
        result.update(run)
    return result


def add_chapter(doc: Document, part: int, chapter: int, path: Path, is_first_in_part: bool = False) -> None:
    paragraphs = read_odt(path)
    title, extras, epigraphs, notes = classify_chapter(paragraphs)
    rel = str(path.relative_to(ROOT))
    for note in notes:
        LOG.exceptions.append(f'- `{rel}`: {note}')
    if title is not None:
        LOG.exceptions.append(f'- `{rel}`: заголовок распознан как абзац {paragraphs[title].number}: “{paragraphs[title].text.strip()}”.')
    else:
        LOG.exceptions.append(f'- `{rel}`: заголовок не распознан; файл начат с новой страницы без добавления нового текста.')

    if not is_first_in_part:
        if title is None or title > 0:
            doc.add_page_break()

    list_indices = is_list_run(list(enumerate(paragraphs)))
    for i, para in enumerate(paragraphs):
        original = para.text
        if not original.strip():
            LOG.change(para.source, para.number, 'blank-layout-paragraph-removed', original, '')
            continue
        runs = normalize_runs(para)
        value = text_of(runs)
        stripped = value.strip()
        if i == title:
            p = doc.add_paragraph(style='Heading 2')
            add_runs(p, runs)
            if title > 0:
                p.paragraph_format.page_break_before = False
            continue
        if i in epigraphs:
            add_normalized_paragraph(doc, para, 'Opening Epigraph', runs)
            continue
        if i in extras:
            add_normalized_paragraph(doc, para, 'Chapter Subtitle' if (title is not None and i > title) else 'Front Note', runs)
            continue
        if stripped.startswith('***'):
            after = stripped[3:].lstrip()
            sep = doc.add_paragraph('***', style='Topic Separator')
            for run in sep.runs:
                set_font(run)
            LOG.change(para.source, para.number, 'topic-separator-separated', stripped, '***' + ('↵' + after if after else ''))
            if after:
                first_style = runs[0][1] if runs else CharStyle()
                add_normalized_paragraph(doc, para, 'Normal', [(after, first_style)])
            continue
        if i in list_indices and re.match(r'^-\s+', stripped):
            add_normalized_paragraph(doc, para, 'Dash List', runs)
            continue
        if i in list_indices and re.match(r'^\d+[.)]\s+', stripped):
            # Preserve text but use the real ordered-list marker instead of a typed marker.
            content = re.sub(r'^\d+[.)]\s+', '', stripped)
            sty = runs[0][1] if runs else CharStyle()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            add_runs(p, [(content, sty)])
            LOG.change(para.source, para.number, 'ordered-list-marker', stripped, content)
            continue
        # Inline/end quotations require authorial review unless their source formatting is unambiguous.
        if '\n' in value:
            LOG.exceptions.append(f'- `{rel}`, абзац {para.number}: сохранён ручной перенос строки; проверьте, является ли он смысловым.')
        if has_initial_bold(runs):
            # Bold-led glossary entries stay ordinary body paragraphs by approved rule.
            add_normalized_paragraph(doc, para, 'Normal', runs)
            continue
        add_normalized_paragraph(doc, para, 'Normal', runs)


def add_front_matter(doc: Document) -> None:
    cover_path = ROOT / '0-01 Обложка.odt'
    if cover_path.exists():
        for para in read_odt(cover_path):
            if para.text.strip():
                add_normalized_paragraph(doc, para, 'Cover Text')
        doc.add_page_break()
    epigraph_path = ROOT / '0-02 Эпиграфы.odt'
    if epigraph_path.exists():
        for para in read_odt(epigraph_path):
            if para.text.strip():
                add_normalized_paragraph(doc, para, 'Opening Epigraph')
        doc.add_page_break()


def main() -> None:
    global ROOT, OUT, CHANGELOG, EXCEPTIONS

    parser = argparse.ArgumentParser(description='Build manuscript docx from ODT chapter files.')
    parser.add_argument('--root', type=Path, default=ROOT, help='Path to book root directory containing ODT files.')
    parser.add_argument('--out', type=Path, default=OUT, help='Output docx file path.')
    args = parser.parse_args()

    ROOT = args.root.resolve()
    OUT = args.out.resolve()
    CHANGELOG = ROOT / 'формат для печати' / 'журнал форматных изменений.csv'
    EXCEPTIONS = ROOT / 'формат для печати' / 'исключения для проверки.md'

    doc = Document()
    configure_section(doc.sections[0])
    # Mirror margins allow left/right margins to become inner/outer on facing pages.
    settings = doc.settings.element
    settings.append(OxmlElement('w:mirrorMargins'))
    define_styles(doc)
    # Front matter deliberately carries no page number in this first working edition.
    add_front_matter(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    body_section.footer.is_linked_to_previous = False
    add_page_number(body_section.footer)
    set_page_start(body_section, 1)
    add_static_toc(doc, toc_entries())

    current_part = None
    for part, chapter, path in document_files():
        if part != current_part:
            section = doc.add_section(WD_SECTION.ODD_PAGE)
            configure_section(section)
            continue_page_numbers(section)
            section.footer.is_linked_to_previous = True
            heading = doc.add_paragraph(PART_TITLES.get(part, f'{part}. ЧАСТЬ'), style='Heading 1')
            for run in heading.runs:
                set_font(run)
            current_part = part
            is_first_in_part = True
        else:
            is_first_in_part = False
        add_chapter(doc, part, chapter, path, is_first_in_part=is_first_in_part)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    with CHANGELOG.open('w', newline='', encoding='utf-8-sig') as fh:
        writer = csv.DictWriter(fh, fieldnames=['source_file', 'paragraph', 'rule', 'old_value', 'new_value'])
        writer.writeheader(); writer.writerows(LOG.rows)
    with EXCEPTIONS.open('w', encoding='utf-8') as fh:
        fh.write('# Исключения и решения для проверки\n\n')
        fh.write('Этот список не меняет текст; он фиксирует места, где автоматическая классификация требует редакторского просмотра.\n\n')
        fh.write('\n'.join(LOG.exceptions) or '- Исключений не обнаружено.')
        fh.write('\n')
    print(f'Created: {OUT}')
    print(f'Changes: {len(LOG.rows)}; exceptions: {len(LOG.exceptions)}')


if __name__ == '__main__':
    main()
